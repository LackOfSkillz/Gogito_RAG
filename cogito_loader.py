# cogito_loader.py
# Cogito Knowledge Loader — Atlas v0.7.1
# - Incremental ingestion via fingerprints (db_index.json)
# - Recursive directory scan of /data
# - Parses PDF/HTML/TXT/MD/DOCX/CSV + OCR for JPG/PNG
# - Splitting + embeddings → Chroma persistent DB (/chroma_db)
# - Progress bars (tqdm) with --progress flag
# - Debug timings with --debug flag
#
# Run:
#   python cogito_loader.py --progress
#   python cogito_loader.py --progress --debug
#   python cogito_loader.py --reset --progress

import os
import re
import io
import sys
import json
import time
import argparse
import hashlib
import datetime as dt
from pathlib import Path
from typing import List, Dict, Tuple

# ---- Third-party deps (parsers / OCR) ----
import fitz  # PyMuPDF for PDFs
from PIL import Image
import pytesseract
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document as DocxDocument

# ---- LangChain v0.2+ compatible imports ----
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

# ---- Progress bars ----
from tqdm.auto import tqdm

# -----------------------
# Configuration
# -----------------------
ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "data"
DB_DIR = ROOT / "chroma_db"
INDEX_PATH = ROOT / "db_index.json"

# Allowed file types
TEXT_EXTS = {".txt", ".md", ".csv"}
DOC_EXTS = {".docx"}
HTML_EXTS = {".html", ".htm"}
PDF_EXTS = {".pdf"}
IMG_EXTS = {".png", ".jpg", ".jpeg"}

ALL_EXTS = TEXT_EXTS | DOC_EXTS | HTML_EXTS | PDF_EXTS | IMG_EXTS

# Splitter defaults (tuned for manuals)
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 200

# Chroma batch sizes
EMBED_BATCH = 256
UPSERT_BATCH = 500

# -----------------------
# Utilities
# -----------------------
def ensure_structure() -> None:
    DATA_DIR.mkdir(exist_ok=True)
    DB_DIR.mkdir(exist_ok=True)
    if not INDEX_PATH.exists():
        INDEX_PATH.write_text(json.dumps({"files": {}, "schema": "atlas-0.7"}, indent=2))

def load_index() -> Dict:
    if INDEX_PATH.exists():
        try:
            return json.loads(INDEX_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {"files": {}, "schema": "atlas-0.7"}

def save_index(idx: Dict) -> None:
    INDEX_PATH.write_text(json.dumps(idx, indent=2))

def file_fingerprint(p: Path) -> str:
    try:
        stat = p.stat()
        sig = f"{p.as_posix()}|{stat.st_size}|{int(stat.st_mtime)}"
    except Exception:
        sig = f"{p.as_posix()}|0|0"
    return hashlib.sha1(sig.encode("utf-8")).hexdigest()

def walk_files(base: Path) -> List[Path]:
    files = []
    for root, _, filenames in os.walk(base):
        for name in filenames:
            fpath = Path(root) / name
            if fpath.suffix.lower() in ALL_EXTS:
                files.append(fpath)
    return files

def clean_text(txt: str) -> str:
    # Basic normalization to help retrieval
    txt = txt.replace("\u00a0", " ")
    txt = re.sub(r"[ \t]+", " ", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    return txt.strip()

# -----------------------
# Parsers
# -----------------------
def parse_pdf(path: Path) -> List[Document]:
    docs = []
    try:
        with fitz.open(path.as_posix()) as pdf:
            for i, page in enumerate(pdf):
                text = page.get_text("text") or ""
                text = clean_text(text)
                if text:
                    docs.append(Document(page_content=text, metadata={
                        "source": path.as_posix(),
                        "type": "pdf",
                        "page": i + 1,
                        "title": path.name
                    }))
    except Exception as e:
        print(f"⚠️  PDF parse failed: {path.name} — {e}")
    return docs

def parse_html(path: Path) -> List[Document]:
    try:
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = clean_text(text)
        if text:
            return [Document(page_content=text, metadata={
                "source": path.as_posix(),
                "type": "html",
                "title": path.name
            })]
    except Exception as e:
        print(f"⚠️  HTML parse failed: {path.name} — {e}")
    return []

def parse_text_like(path: Path) -> List[Document]:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        text = clean_text(text)
        if text:
            return [Document(page_content=text, metadata={
                "source": path.as_posix(),
                "type": path.suffix.lower().lstrip("."),
                "title": path.name
            })]
    except Exception as e:
        print(f"⚠️  Text parse failed: {path.name} — {e}")
    return []

def parse_docx(path: Path) -> List[Document]:
    try:
        doc = DocxDocument(path.as_posix())
        parts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        text = clean_text("\n".join(parts))
        if text:
            return [Document(page_content=text, metadata={
                "source": path.as_posix(),
                "type": "docx",
                "title": path.name
            })]
    except Exception as e:
        print(f"⚠️  DOCX parse failed: {path.name} — {e}")
    return []

def parse_csv(path: Path) -> List[Document]:
    try:
        # Limit rows in memory representation; CSV becomes text block
        df = pd.read_csv(path.as_posix(), nrows=100000)  # safety cap
        text = df.to_csv(index=False)
        text = clean_text(text)
        if text:
            return [Document(page_content=text, metadata={
                "source": path.as_posix(),
                "type": "csv",
                "title": path.name
            })]
    except Exception as e:
        print(f"⚠️  CSV parse failed: {path.name} — {e}")
    return []

def ocr_image(path: Path) -> List[Document]:
    docs = []
    try:
        img = Image.open(path.as_posix())
        text = pytesseract.image_to_string(img) or ""
        text = clean_text(text)
        if text:
            docs.append(Document(page_content=text, metadata={
                "source": path.as_posix(),
                "type": "image_ocr",
                "title": path.name
            }))
    except Exception as e:
        print(f"⚠️  OCR failed: {path.name} — {e}")
    return docs

def parse_file(path: Path) -> List[Document]:
    ext = path.suffix.lower()
    if ext in PDF_EXTS:
        return parse_pdf(path)
    if ext in HTML_EXTS:
        return parse_html(path)
    if ext in TEXT_EXTS and ext != ".csv":
        return parse_text_like(path)
    if ext == ".csv":
        return parse_csv(path)
    if ext in DOC_EXTS:
        return parse_docx(path)
    if ext in IMG_EXTS:
        return ocr_image(path)
    return []

# -----------------------
# Embedding & DB
# -----------------------
def get_splitter():
    return RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n## ", "\n# ", "\n\n", "\n", " ", ""],
    )

def get_embeddings():
    # all-MiniLM-L6-v2 is fast; upgrade later if needed
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

def get_db(embeddings):
    return Chroma(persist_directory=str(DB_DIR), embedding_function=embeddings)

# -----------------------
# Main
# -----------------------
def main():
    parser = argparse.ArgumentParser(description="Cogito RAG Loader — Atlas v0.7.1")
    parser.add_argument("--progress", action="store_true", help="Show live progress bars")
    parser.add_argument("--debug", action="store_true", help="Verbose stage timings")
    parser.add_argument("--reset", action="store_true", help="Delete DB and re-ingest everything")
    parser.add_argument("--limit", type=int, default=0, help="Limit number of files (for testing)")
    args = parser.parse_args()

    SHOW_PROGRESS = args.progress or os.getenv("COGITO_PROGRESS", "0") == "1"
    DEBUG = args.debug
    RESET = args.reset

    print("⚙️ Starting ingestion run (Atlas v0.7.1)")
    t0 = time.time()

    ensure_structure()

    # Reset DB if requested
    if RESET and DB_DIR.exists():
        for item in DB_DIR.glob("**/*"):
            try:
                if item.is_file():
                    item.unlink()
            except Exception:
                pass
        for sub in sorted(DB_DIR.glob("**/*"), reverse=True):
            try:
                if sub.is_dir():
                    sub.rmdir()
            except Exception:
                pass
        DB_DIR.mkdir(exist_ok=True)
        print("🧹 Database folder cleared.")

    # Load previous index & scan files
    idx = load_index()
    all_paths = walk_files(DATA_DIR)

    if args.limit > 0:
        all_paths = all_paths[: args.limit]

    total_files = len(all_paths)
    print(f"Found {total_files} source files to process")

    # Fingerprint scan
    new_or_modified: List[Path] = []
    removed: List[str] = []
    unchanged: List[Path] = []

    prev_map: Dict[str, str] = idx.get("files", {})

    # Hashing with progress
    iter_hash = tqdm(all_paths, desc="🔑 Hashing", unit="file", disable=not SHOW_PROGRESS)
    current_map = {}
    for p in iter_hash:
        fp = file_fingerprint(p)
        current_map[p.as_posix()] = fp
        if prev_map.get(p.as_posix()) != fp:
            new_or_modified.append(p)
        else:
            unchanged.append(p)

    # Detect removed
    prev_paths = set(prev_map.keys())
    cur_paths = set(current_map.keys())
    for gone in prev_paths - cur_paths:
        removed.append(gone)

    print(f"\n📄 {len(new_or_modified)} new or modified file(s).")
    print(f"🗑️  {len(removed)} removed file(s).")
    print(f"⏭️  {len(unchanged)} unchanged file(s) skipped.\n")

    # Quick exit if nothing to do
    if not new_or_modified and not removed:
        print("✅ Nothing to ingest. Database is up to date.")
        return

    # Parse new/modified files
    parse_iter = tqdm(new_or_modified, desc="📄 Parsing", unit="file", disable=not SHOW_PROGRESS)
    parsed_docs: List[Document] = []
    n_ocr = 0
    for path in parse_iter:
        docs = parse_file(path)
        # Track OCR count
        if path.suffix.lower() in IMG_EXTS:
            n_ocr += 1 if docs else 0
        parsed_docs.extend(docs)

    # Split into chunks
    splitter = get_splitter()
    split_iter = tqdm(parsed_docs, desc="✂️ Splitting", unit="doc", disable=not SHOW_PROGRESS)
    chunks: List[Document] = []
    for d in split_iter:
        # Use split_documents to preserve metadata per chunk
        parts = splitter.split_documents([d])
        chunks.extend(parts)

    # Initialize embeddings & DB
    if DEBUG:
        print("⚙️ Initializing embeddings + DB...")
    t_embed0 = time.time()
    embeddings = get_embeddings()
    db = get_db(embeddings)

    # Upsert chunks in batches (Chroma has max batch sizes internally)
    def batched(seq, size):
        for i in range(0, len(seq), size):
            yield seq[i : i + size]

    added = 0
    batch_iter = tqdm(list(batched(chunks, UPSERT_BATCH)), desc="💾 Upsert", unit="batch", disable=not SHOW_PROGRESS)
    for batch in batch_iter:
        db.add_documents(batch)
        added += len(batch)

    # Persist (Chroma >=0.4 auto-persists, but calling is harmless)
    try:
        db.persist()
    except Exception:
        pass

    # Update index: remove deleted paths, store current hashes
    files_map = idx.get("files", {})
    for gone in removed:
        files_map.pop(gone, None)
    for k, v in current_map.items():
        files_map[k] = v
    idx["files"] = files_map
    save_index(idx)

    # Stats
    t1 = time.time()
    elapsed = t1 - t0
    docs_count = len(parsed_docs)
    print("\n================= Summary =================")
    print(f"🖼️ OCR processed         : ~{n_ocr} image(s)")
    print(f"📘 Parsed documents      : {docs_count}")
    print(f"✂️  Chunks created        : {len(chunks)}")
    print(f"💾 Chunks upserted       : {added}")
    print(f"⏱  Total time            : {elapsed:.2f}s")
    if docs_count:
        print(f"⚡ Files/sec (parse)      : {len(new_or_modified)/max(1.0, elapsed):.2f}")
    print(f"✅ Incremental update complete. Database saved to: {DB_DIR.as_posix()}")
    print("===========================================\n")

if __name__ == "__main__":
    main()
